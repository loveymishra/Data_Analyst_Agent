import pandas as pd
import PyPDF2
from docx import Document
from PIL import Image
import io
import streamlit as st

def parse_csv(file):
    """
    Parse CSV file and return pandas DataFrame
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        pandas.DataFrame: Parsed CSV data
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                file.seek(0)
                df = pd.read_csv(file, encoding=encoding)
                return df
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try with errors='ignore'
        file.seek(0)
        df = pd.read_csv(file, encoding='utf-8', errors='ignore')
        return df
        
    except Exception as e:
        raise Exception(f"Error parsing CSV file: {str(e)}")

def parse_excel(file):
    """
    Parse Excel file and return pandas DataFrame
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        pandas.DataFrame: Parsed Excel data
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        # Read Excel file (first sheet by default)
        df = pd.read_excel(file, engine='openpyxl')
        return df
        
    except Exception as e:
        raise Exception(f"Error parsing Excel file: {str(e)}")

def parse_pdf(file):
    """
    Parse PDF file and extract text content
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        # Create PDF reader object
        pdf_reader = PyPDF2.PdfReader(file)
        
        # Extract text from all pages
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        if not text.strip():
            raise Exception("No text could be extracted from the PDF")
            
        return text.strip()
        
    except Exception as e:
        raise Exception(f"Error parsing PDF file: {str(e)}")

def parse_docx(file):
    """
    Parse DOCX file and extract text content
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        # Create Document object
        doc = Document(file)
        
        # Extract text from all paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise Exception("No text could be extracted from the DOCX file")
            
        return text.strip()
        
    except Exception as e:
        raise Exception(f"Error parsing DOCX file: {str(e)}")

def parse_image(file):
    """
    Parse image file and return PIL Image object
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        PIL.Image: Loaded image object
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        # Open image with PIL
        image = Image.open(file)
        
        # Convert to RGB if necessary (for consistency)
        if image.mode in ('RGBA', 'LA', 'P'):
            image = image.convert('RGB')
            
        return image
        
    except Exception as e:
        raise Exception(f"Error parsing image file: {str(e)}")

# Utility functions for data validation and preprocessing

def validate_dataframe(df):
    """
    Validate and clean DataFrame
    
    Args:
        df: pandas DataFrame
        
    Returns:
        pandas.DataFrame: Cleaned DataFrame
    """
    try:
        # Remove completely empty rows and columns
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        # Clean column names (remove extra spaces, special characters)
        df.columns = df.columns.str.strip()
        
        # Try to infer better data types
        df = df.infer_objects()
        
        # Convert string representations of numbers to actual numbers
        for col in df.columns:
            if df[col].dtype == 'object':
                # Try to convert to numeric, ignoring errors
                numeric_series = pd.to_numeric(df[col], errors='coerce')
                if not numeric_series.isna().all():
                    df[col] = numeric_series
        
        return df
        
    except Exception as e:
        print(f"Warning: Error during DataFrame validation: {str(e)}")
        return df

def get_file_info(file):
    """
    Get basic information about uploaded file
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        dict: File information
    """
    try:
        info = {
            'name': file.name,
            'size': file.size,
            'type': file.type,
            'extension': file.name.split('.')[-1].lower() if '.' in file.name else 'unknown'
        }
        return info
    except Exception as e:
        return {'error': str(e)}

def preview_text(text, max_length=500):
    """
    Create a preview of text content
    
    Args:
        text: Text content
        max_length: Maximum length of preview
        
    Returns:
        str: Preview text
    """
    if len(text) <= max_length:
        return text
    else:
        return text[:max_length] + "..."
